"""Extraction of plain text and provenance metadata from uploaded files.

Pure functions over bytes: no database access, no configuration, no network.
That keeps the format-specific logic — which is where the awkward edge cases
live — independently testable.
"""

import io
import re
from dataclasses import dataclass
from pathlib import Path

import docx
import pypdf
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.core.constants import SUPPORTED_CONTENT_TYPES, SUPPORTED_EXTENSIONS
from app.core.exceptions import (
    DocumentParseError,
    EmptyDocumentError,
    UnsupportedDocumentTypeError,
)

_MARKDOWN_HEADING = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<title>.+?)\s*#*$")
_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1252", "latin-1")

# Resolved once: `qn` builds a Clark-notation tag on every call.
_TEXT_BOX_TAG = qn("w:txbxContent")
_PARAGRAPH_TAG = qn("w:p")
_TEXT_TAG = qn("w:t")

# Markup Compatibility, which `python-docx` does not register a prefix for.
_FALLBACK_TAG = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"


@dataclass(frozen=True)
class ParsedSection:
    """A slice of extracted text together with where it came from.

    `page_number` is 1-indexed and only available for paginated formats.
    `section_title` is the nearest preceding heading, where the format exposes
    one.
    """

    text: str
    page_number: int | None = None
    section_title: str | None = None


@dataclass(frozen=True)
class ParsedDocument:
    """The full result of parsing one uploaded file."""

    sections: list[ParsedSection]
    page_count: int | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(section.text for section in self.sections)


def resolve_format(filename: str, content_type: str | None) -> str:
    """Return the short format key for an upload, or raise.

    The declared content type is trusted first, since browsers set it
    reliably for these formats; the extension is a fallback for clients that
    send `application/octet-stream`.
    """

    normalized_type = (content_type or "").split(";")[0].strip().lower()
    if normalized_type in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[normalized_type]

    extension = Path(filename).suffix.lower()
    if extension in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[extension]

    raise UnsupportedDocumentTypeError(
        f"Unsupported document type: {content_type or extension or 'unknown'}. "
        f"Supported formats are PDF, DOCX, TXT and Markdown."
    )


def _decode(data: bytes) -> str:
    for encoding in _TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentParseError("File could not be decoded as text.")


def _parse_pdf(data: bytes) -> ParsedDocument:
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        page_count = len(reader.pages)
        sections = [
            ParsedSection(text=text.strip(), page_number=number)
            for number, page in enumerate(reader.pages, start=1)
            if (text := page.extract_text() or "").strip()
        ]
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"PDF could not be read: {exc}") from exc

    return ParsedDocument(sections=sections, page_count=page_count)


def _iter_body_blocks(document: docx.document.Document):
    """Yield paragraphs and tables in the order they appear in the document.

    `document.paragraphs` and `document.tables` are two flat lists with no
    interleaving, so reading them separately would append every table after
    every paragraph. Walking the body's own children is the only way to keep a
    table sitting between the paragraphs it belongs to.
    """

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _text_box_blocks(element) -> list[str]:
    """Return the text of every text box anchored inside `element`.

    Text box content lives in `w:txbxContent`, nested under `w:drawing`
    (DrawingML) or `w:pict` (legacy VML). Either way the marker element is the
    same, so one search catches both. `python-docx` does not surface this at
    all: `Paragraph.text` reads `w:t` from direct runs, and a text box's runs
    sit several levels deeper — which is why documents whose content is laid
    out entirely in shapes currently extract as empty.

    Word usually writes both encodings of the same shape, wrapped in an
    `mc:AlternateContent` pair: `mc:Choice` holds the DrawingML a modern
    reader uses, `mc:Fallback` the VML an old one does. They carry identical
    text, so the fallback is skipped or every shape would be read twice.
    """

    blocks: list[str] = []

    # `iter` with qualified tags rather than XPath: `w:txbxContent` has no
    # registered `python-docx` element class, so it comes back as a plain lxml
    # element whose `.xpath()` has no namespace map — while its `w:p` children
    # do have one and reject a `namespaces=` argument. `iter` sidesteps the
    # inconsistency and skips the XPath engine entirely.
    for box in element.iter(_TEXT_BOX_TAG):
        # A text box inside another is already covered by the outer one's
        # descendant walk; taking both would duplicate its text.
        if any(
            ancestor.tag in (_TEXT_BOX_TAG, _FALLBACK_TAG)
            for ancestor in box.iterancestors()
        ):
            continue

        lines = [
            line
            for paragraph in box.iter(_PARAGRAPH_TAG)
            if (
                line := "".join(
                    node.text or "" for node in paragraph.iter(_TEXT_TAG)
                ).strip()
            )
        ]

        if body := "\n".join(lines).strip():
            blocks.append(body)

    return blocks


def _cell_text(cell: _Cell) -> str:
    """Flatten one table cell to a single line.

    Includes nested tables and text boxes, both of which `_Cell.text` skips —
    it reads only the cell's direct paragraphs.
    """

    parts = [" ".join(cell.text.split())]
    parts.extend(_text_box_blocks(cell._tc))
    parts.extend(
        " ".join(_cell_text(inner) for inner in row.cells)
        for nested in cell.tables
        for row in nested.rows
    )

    return " ".join(part for part in parts if part.strip()).strip()


def _table_rows(table: Table) -> list[list[str]]:
    """Return non-empty rows as lists of cell text, merged cells collapsed."""

    rows: list[list[str]] = []

    for row in table.rows:
        cells: list[str] = []
        seen: set[int] = set()

        for cell in row.cells:
            # A horizontally merged cell is returned once per column it spans,
            # backed by the same `w:tc`. Emitting it repeatedly would restate
            # its value across several columns.
            if id(cell._tc) in seen:
                continue

            seen.add(id(cell._tc))
            cells.append(_cell_text(cell))

        if any(cells):
            rows.append(cells)

    return rows


def _serialise_table(table: Table) -> str:
    """Render a table as one `Header: value` line per row.

    Chosen over a Markdown grid because a chunk boundary can fall anywhere:
    a grid row separated from its header line becomes unreadable, whereas a
    header-qualified row still says what each value means. It also embeds
    closer to prose than pipe-delimited columns, and reads naturally when the
    model receives it as context.

    Tables with no usable header — a single row, a single column, or a first
    row that does not look like labels — fall back to plain delimited rows,
    which is the right shape for the layout tables DOCX authors use for
    formatting rather than data.
    """

    try:
        rows = _table_rows(table)
    except Exception:
        # Malformed grids (bad gridSpan, truncated rows) should cost their
        # own content, not the whole document.
        return ""

    if not rows:
        return ""

    header = rows[0]
    labels = [cell for cell in header if cell]
    has_header = len(rows) > 1 and len(labels) >= 2 and len(set(labels)) == len(labels)

    if not has_header:
        return "\n".join(" | ".join(cell for cell in row if cell) for row in rows)

    lines: list[str] = []

    for row in rows[1:]:
        pairs = []
        for position, value in enumerate(row):
            if not value:
                continue

            label = header[position] if position < len(header) else ""
            pairs.append(f"{label}: {value}" if label else value)

        if pairs:
            lines.append(" | ".join(pairs))

    # A header with no data rows beneath it is still worth keeping.
    return "\n".join(lines) if lines else " | ".join(labels)


def _parse_docx(data: bytes) -> ParsedDocument:
    """Extract DOCX paragraphs, tables and text boxes, in document order.

    DOCX has no reliable page boundaries without rendering the document, so
    `page_number` stays None and headings carry the provenance instead.

    Tables and text boxes each become their own section so that a chunk never
    mixes a table row with surrounding prose — the same invariant the page and
    heading splits already maintain.
    """

    try:
        document = docx.Document(io.BytesIO(data))
        blocks = list(_iter_body_blocks(document))
    except Exception as exc:
        raise DocumentParseError(f"DOCX could not be read: {exc}") from exc

    sections: list[ParsedSection] = []
    current_title: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            sections.append(ParsedSection(text=body, section_title=current_title))
        buffer.clear()

    def emit(text: str) -> None:
        flush()
        sections.append(ParsedSection(text=text, section_title=current_title))

    for block in blocks:
        if isinstance(block, Table):
            if serialised := _serialise_table(block):
                emit(serialised)
            continue

        text = block.text.strip()

        if text:
            style_name = (block.style.name or "") if block.style else ""
            if style_name.lower().startswith("heading") or style_name == "Title":
                flush()
                current_title = text
            else:
                buffer.append(text)

        # Runs after the paragraph's own text, so a shape anchored to a
        # paragraph lands after it rather than before.
        for box in _text_box_blocks(block._p):
            emit(box)

    flush()
    return ParsedDocument(sections=sections, page_count=None)


def _parse_markdown(data: bytes) -> ParsedDocument:
    """Split Markdown on ATX headings so each section keeps its title."""

    content = _decode(data)
    sections: list[ParsedSection] = []
    current_title: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            sections.append(ParsedSection(text=body, section_title=current_title))
        buffer.clear()

    for line in content.splitlines():
        match = _MARKDOWN_HEADING.match(line.strip())
        if match:
            flush()
            current_title = match.group("title").strip()
            continue

        buffer.append(line)

    flush()
    return ParsedDocument(sections=sections, page_count=None)


def _parse_text(data: bytes) -> ParsedDocument:
    content = _decode(data).strip()
    sections = [ParsedSection(text=content)] if content else []
    return ParsedDocument(sections=sections, page_count=None)


_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "md": _parse_markdown,
    "txt": _parse_text,
}


def parse_document(
    data: bytes, filename: str, content_type: str | None
) -> ParsedDocument:
    """Extract text and provenance from an uploaded file.

    Raises `UnsupportedDocumentTypeError` for unknown formats,
    `DocumentParseError` when a file of a known format cannot be read, and
    `EmptyDocumentError` when it yields no text — the last of which covers
    scanned PDFs with no text layer.
    """

    if not data:
        raise EmptyDocumentError("Uploaded file is empty.")

    document_format = resolve_format(filename, content_type)
    parsed = _PARSERS[document_format](data)

    if not parsed.sections:
        raise EmptyDocumentError(
            "No text could be extracted from the document. Scanned documents "
            "without a text layer are not supported."
        )

    return parsed
