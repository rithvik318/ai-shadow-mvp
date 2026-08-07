"""Builders for realistic upload fixtures, generated in memory.

Binary fixtures are built rather than committed so the test suite has no
opaque blobs, and so page and heading structure is visible in the test that
depends on it.
"""

import io

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# `python-docx` registers neither prefix, so both are declared by hand.
_MC_NS = 'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
_WPS_NS = (
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
)


def build_pdf(pages: list[str]) -> bytes:
    """Return a minimal, valid single-font PDF with one text line per page.

    Written by hand rather than with a PDF generation library so the test
    suite needs no dependency that production code does not already have.
    """

    objects: list[bytes] = []

    def add(body: bytes) -> int:
        objects.append(body)
        return len(objects)

    font_id = 3 + 2 * len(pages)
    page_ids = [3 + 2 * index for index in range(len(pages))]

    # Placeholder ordering: catalog (1), pages (2), then page/content pairs.
    add(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    add(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    for index, text in enumerate(pages):
        content_id = page_ids[index] + 1
        add(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
                f"/Contents {content_id} 0 R >>"
            ).encode()
        )
        escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        add(b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream))

    add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []

    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % number + body + b"\nendobj\n"

    xref_offset = len(out)
    out += b"xref\n0 %d\n" % (len(objects) + 1)
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset

    out += b"trailer\n<< /Size %d /Root 1 0 R >>\n" % (len(objects) + 1)
    out += b"startxref\n%d\n%%%%EOF\n" % xref_offset

    return bytes(out)


def build_docx(sections: list[tuple[str, str]]) -> bytes:
    """Return a DOCX built from (heading, body) pairs."""

    document = docx.Document()
    for heading, body in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_markdown(sections: list[tuple[str, str]]) -> bytes:
    """Return Markdown bytes built from (heading, body) pairs."""

    parts = [f"# {heading}\n\n{body}\n" for heading, body in sections]
    return "\n".join(parts).encode("utf-8")


def build_text(body: str) -> bytes:
    return body.encode("utf-8")


def build_docx_blocks(blocks: list[tuple[str, object]]) -> bytes:
    """Return a DOCX assembled from ordered blocks.

    Each block is a (kind, payload) pair:

        ("heading",   "Overview")
        ("paragraph", "Body text.")
        ("table",     [["Name", "Role"], ["Ann", "Lead"]])
        ("textbox",   "Text inside a shape")
        ("merged",    [["Name", "Role"], ["Ann", "Lead"]])   # row 0 merged

    Written in the order given, so a test can assert that the parser preserves
    document order rather than grouping by type. `python-docx` has no API for
    text boxes, so those are injected as raw `w:txbxContent` — the same element
    Word writes for both DrawingML and legacy VML shapes.
    """

    document = docx.Document()

    for kind, payload in blocks:
        if kind == "heading":
            document.add_heading(str(payload), level=1)
        elif kind == "paragraph":
            document.add_paragraph(str(payload))
        elif kind == "textbox":
            paragraph = document.add_paragraph()
            lines = "".join(
                f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>"
                for line in str(payload).split("\n")
            )
            paragraph._p.append(
                parse_xml(
                    f"<w:r {nsdecls('w')}><w:pict><w:txbxContent>"
                    f"{lines}"
                    f"</w:txbxContent></w:pict></w:r>"
                )
            )
        elif kind in ("table", "merged"):
            grid = payload
            table = document.add_table(rows=len(grid), cols=len(grid[0]))
            for r, row in enumerate(grid):
                for c, value in enumerate(row):
                    table.cell(r, c).text = value
            if kind == "merged":
                table.cell(0, 0).merge(table.cell(0, 1))
        else:
            raise ValueError(f"unknown block kind: {kind}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx_alternate_content_text_box(text: str) -> bytes:
    """Return a DOCX whose text box is written in both shape encodings.

    This is what Word itself emits: an `mc:AlternateContent` pair holding the
    same text as DrawingML (`mc:Choice`) and as legacy VML (`mc:Fallback`).
    Only the DrawingML branch is displayed, so a reader that takes both sees
    every shape twice.
    """

    document = docx.Document()
    paragraph = document.add_paragraph()
    body = f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
    paragraph._p.append(
        parse_xml(
            f"<w:r {nsdecls('w')} {_MC_NS} {_WPS_NS}>"
            f'<mc:AlternateContent><mc:Choice Requires="wps">'
            f"<w:drawing><wps:txbx><w:txbxContent>{body}</w:txbxContent>"
            f"</wps:txbx></w:drawing>"
            f"</mc:Choice><mc:Fallback>"
            f"<w:pict><w:txbxContent>{body}</w:txbxContent></w:pict>"
            f"</mc:Fallback></mc:AlternateContent></w:r>"
        )
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx_table_with_text_box(grid: list[list[str]], text: str) -> bytes:
    """Return a DOCX table whose last cell also contains a text box.

    Word's capability-statement layouts routinely put a shape inside a table
    cell, where `_Cell.text` cannot see it.
    """

    document = docx.Document()
    table = document.add_table(rows=len(grid), cols=len(grid[0]))
    for r, row in enumerate(grid):
        for c, value in enumerate(row):
            table.cell(r, c).text = value

    cell = table.cell(len(grid) - 1, len(grid[0]) - 1)
    cell.paragraphs[0]._p.append(
        parse_xml(
            f"<w:r {nsdecls('w')}><w:pict><w:txbxContent>"
            f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
            f"</w:txbxContent></w:pict></w:r>"
        )
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx_nested_table(outer: str, inner: list[list[str]]) -> bytes:
    """Return a DOCX whose table cell contains another table."""

    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = outer

    nested = cell.add_table(rows=len(inner), cols=len(inner[0]))
    for r, row in enumerate(inner):
        for c, value in enumerate(row):
            nested.cell(r, c).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
